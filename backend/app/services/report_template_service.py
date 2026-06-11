from html import escape
from typing import Any

from docx import Document

from backend.app.models.report import AiReport


class ReportTemplateService:
    def build_template_data(self, report: AiReport) -> dict[str, Any]:
        content = report.content_json or {}
        return {
            "title": content.get("title") or report.title,
            "summary": content.get("summary") or "",
            "sections": self._normalize_sections(content.get("sections", [])),
            "risks": self._normalize_text_list(content.get("risks", [])),
            "recommendations": self._normalize_text_list(content.get("recommendations", [])),
            "source_refs": self._normalize_text_list(content.get("source_refs", [])),
        }

    def write_docx(self, report: AiReport, file_path) -> None:
        data = self.build_template_data(report)
        document = Document()
        document.add_heading(data["title"], 0)
        if data["summary"]:
            document.add_heading("报告摘要", level=1)
            document.add_paragraph(data["summary"])
        for section in data["sections"]:
            document.add_heading(section["heading"], level=1)
            if section["content"]:
                document.add_paragraph(section["content"])
            for metric in section["metrics"]:
                document.add_paragraph(f"{metric['name']}：{metric['value']}", style="List Bullet")
        self._add_list_section(document, "风险提示", data["risks"])
        self._add_list_section(document, "建议事项", data["recommendations"])
        self._add_list_section(document, "来源引用", data["source_refs"])
        document.save(file_path)

    def render_html(self, report: AiReport) -> str:
        data = self.build_template_data(report)
        sections_html = "".join(self._section_html(section) for section in data["sections"])
        risks_html = self._list_html("风险提示", data["risks"])
        recommendations_html = self._list_html("建议事项", data["recommendations"])
        source_refs_html = self._list_html("来源引用", data["source_refs"])
        return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <style>
    body {{ font-family: SimSun, serif; color: #1f2933; line-height: 1.6; }}
    h1 {{ font-size: 24px; text-align: center; margin-bottom: 18px; }}
    h2 {{ font-size: 18px; margin-top: 20px; border-bottom: 1px solid #d9e2ec; padding-bottom: 4px; }}
    p {{ font-size: 12px; }}
    table {{ width: 100%; border-collapse: collapse; margin-top: 8px; }}
    th, td {{ border: 1px solid #bcccdc; padding: 6px; font-size: 11px; }}
    th {{ background: #f0f4f8; }}
  </style>
</head>
<body>
  <h1>{escape(str(data["title"]))}</h1>
  <h2>报告摘要</h2>
  <p>{escape(str(data["summary"]))}</p>
  {sections_html}
  {risks_html}
  {recommendations_html}
  {source_refs_html}
</body>
</html>"""

    def _normalize_sections(self, sections: Any) -> list[dict[str, Any]]:
        if not isinstance(sections, list):
            return []
        normalized = []
        for section in sections:
            if not isinstance(section, dict):
                continue
            normalized.append(
                {
                    "heading": str(section.get("heading") or ""),
                    "content": str(section.get("content") or ""),
                    "metrics": self._normalize_metrics(section.get("metrics", [])),
                }
            )
        return normalized

    def _normalize_metrics(self, metrics: Any) -> list[dict[str, str]]:
        if not isinstance(metrics, list):
            return []
        normalized = []
        for metric in metrics:
            if not isinstance(metric, dict):
                continue
            normalized.append(
                {
                    "name": str(metric.get("name") or ""),
                    "value": str(metric.get("value") if metric.get("value") is not None else ""),
                }
            )
        return normalized

    def _normalize_text_list(self, values: Any) -> list[str]:
        if not isinstance(values, list):
            return []
        return [str(value) for value in values if value is not None]

    def _add_list_section(self, document: Document, title: str, values: list[str]) -> None:
        if not values:
            return
        document.add_heading(title, level=1)
        for value in values:
            document.add_paragraph(value, style="List Bullet")

    def _section_html(self, section: dict[str, Any]) -> str:
        metrics_rows = "".join(
            f"<tr><td>{escape(metric['name'])}</td><td>{escape(metric['value'])}</td></tr>"
            for metric in section["metrics"]
        )
        metrics_html = f"<table><tr><th>指标</th><th>数值</th></tr>{metrics_rows}</table>" if metrics_rows else ""
        return (
            f"<h2>{escape(section['heading'])}</h2>"
            f"<p>{escape(section['content'])}</p>"
            f"{metrics_html}"
        )

    def _list_html(self, title: str, values: list[str]) -> str:
        if not values:
            return ""
        items = "".join(f"<li>{escape(value)}</li>" for value in values)
        return f"<h2>{escape(title)}</h2><ul>{items}</ul>"
